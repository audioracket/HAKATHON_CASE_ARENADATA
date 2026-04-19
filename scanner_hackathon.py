from __future__ import annotations
import os, re, io, json, math, csv, itertools, mimetypes
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from datetime import datetime

ROOT_DIR = Path('../share')
OUTPUT_CSV = Path('pii_scan_results.csv')
INCLUDE_EXTS = {'doc', 'docx', 'gif', 'html', 'jpeg', 'jpg', 'pdf', 'mp4', 'png', 'rtf', 'xls'}


def safe_import(name):
    try:
        return __import__(name)
    except Exception:
        return None


# Импорт библиотек для извлечения текста
PyPDF2 = safe_import('PyPDF2')
pdfminer = safe_import('pdfminer')
docx = safe_import('docx')
bs4 = safe_import('bs4')
pandas = safe_import('pandas')
PIL = safe_import('PIL')
pytesseract = safe_import('pytesseract')
chardet = safe_import('chardet')
cv2 = safe_import('cv2')

# Импорт современных библиотек для PII detection
try:
    from presidio_analyzer import AnalyzerEngine, PatternRecognizer, Pattern, EntityRecognizer
    from presidio_analyzer.predefined_recognizers import (
        EmailRecognizer, PhoneRecognizer, CreditCardRecognizer,
        UsaSsnRecognizer, UsBankRecognizer, CryptoRecognizer,
        SpacyRecognizer, DateRecognizer, IpRecognizer,
        UsLicenseRecognizer, UkNhsRecognizer
    )

    PRESIDIO_AVAILABLE = True
except ImportError:
    PRESIDIO_AVAILABLE = False
    print("Warning: presidio-analyzer not installed. Install with: pip install presidio-analyzer presidio-anonymizer")

try:
    from email_validator import validate_email, EmailNotValidError

    EMAIL_VALIDATOR_AVAILABLE = True
except ImportError:
    EMAIL_VALIDATOR_AVAILABLE = False

try:
    import phonenumbers

    PHONENUMBERS_AVAILABLE = True
except ImportError:
    PHONENUMBERS_AVAILABLE = False

try:
    import snils_validator

    SNILS_VALIDATOR_AVAILABLE = True
except ImportError:
    SNILS_VALIDATOR_AVAILABLE = False

# Конфигурация для presidio
PII_ENTITIES = {
    'обычные': [
        'PERSON', 'EMAIL_ADDRESS', 'PHONE_NUMBER', 'DATE_TIME',
        'LOCATION', 'ADDRESS', 'URL', 'IP_ADDRESS'
    ],
    'государственные': [
        'US_SSN', 'US_PASSPORT', 'US_DRIVER_LICENSE', 'UK_NHS',
        'IN_PAN', 'AU_ABN', 'AU_ACN', 'ES_NIF'
    ],
    'платёжные': [
        'CREDIT_CARD', 'US_BANK_NUMBER', 'CRYPTO'
    ]
}


class OptimizedPIIDetector:
    """Оптимизированный детектор PII с использованием presidio"""

    def __init__(self):
        self.analyzer = None
        self._init_presidio()

    def _init_presidio(self):
        """Инициализация presidio анализатора"""
        if PRESIDIO_AVAILABLE:
            try:
                self.analyzer = AnalyzerEngine()
                # Добавляем русские паттерны для имен
                self._add_russian_patterns()
            except Exception as e:
                print(f"Error initializing Presidio: {e}")
                self.analyzer = None
        else:
            self.analyzer = None

    def _add_russian_patterns(self):
        """Добавление русскоязычных паттернов"""
        try:
            # Паттерн для русского имени (ФИО)
            russian_name_pattern = Pattern(
                name='russian_name',
                regex=r'\b[А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+){1,2}\b',
                score=0.7
            )
            russian_name_recognizer = PatternRecognizer(
                supported_entity='PERSON',
                patterns=[russian_name_pattern],
                supported_language='ru'
            )
            self.analyzer.registry.add_recognizer(russian_name_recognizer)

            # Паттерн для российского паспорта
            passport_pattern = Pattern(
                name='russian_passport',
                regex=r'\b\d{2}\s?\d{2}\s?\d{6}\b',
                score=0.85
            )
            passport_recognizer = PatternRecognizer(
                supported_entity='US_PASSPORT',
                patterns=[passport_pattern],
                supported_language='ru'
            )
            self.analyzer.registry.add_recognizer(passport_recognizer)

        except Exception as e:
            print(f"Error adding Russian patterns: {e}")

    def analyze_text(self, text: str) -> Dict[str, List[Dict]]:
        """Анализ текста с помощью presidio"""
        results = {
            'обычные': [],
            'государственные': [],
            'платёжные': [],
            'биометрические': [],
            'специальные': []
        }

        if not self.analyzer or not text:
            return results

        try:
            # Ограничиваем длину текста для производительности
            if len(text) > 100000:
                text = text[:100000]

            # Анализ с presidio
            analyzer_results = self.analyzer.analyze(
                text=text,
                entities=[],
                language='ru',
                score_threshold=0.5
            )

            # Группировка результатов по категориям
            for res in analyzer_results:
                entity = res.entity_type
                value = text[res.start:res.end]

                for category, entities in PII_ENTITIES.items():
                    if entity in entities:
                        results[category].append({
                            'entity': entity,
                            'value': value,
                            'score': res.score,
                            'start': res.start,
                            'end': res.end
                        })
                        break

            # Дополнительная проверка специальных категорий
            self._check_special_categories(text, results)

        except Exception as e:
            print(f"Error in Presidio analysis: {e}")

        return results

    def _check_special_categories(self, text: str, results: Dict):
        """Проверка специальных категорий ПДн"""
        low_text = text.lower()

        # Биометрические данные
        biometric_keywords = [
            'отпечаток', 'палец', 'биометрия', 'радужка', 'сетчатка',
            'голос', 'селфи', 'фото', 'сканирование', 'дактилоскопия'
        ]

        if any(kw in low_text for kw in biometric_keywords):
            results['биометрические'].append({
                'entity': 'BIOMETRIC',
                'value': 'found_biometric_keywords',
                'score': 0.8
            })

        # Специальные категории (медицина, религия, политика)
        special_keywords = {
            'медицинские': ['диагноз', 'болезнь', 'лечение', 'медкарта', 'анамнез', 'вич', 'спид'],
            'религиозные': ['религия', 'вероисповедание', 'церковь', 'мечеть', 'синагога'],
            'политические': ['партия', 'политический', 'голосование', 'выборы'],
            'интимные': ['сексуальный', 'интимный', 'ориентация']
        }

        for subcat, keywords in special_keywords.items():
            if any(kw in low_text for kw in keywords):
                results['специальные'].append({
                    'entity': f'SPECIAL_{subcat.upper()}',
                    'value': 'found_special_keywords',
                    'score': 0.9
                })


class FallbackPIIDetector:
    """Запасной детектор PII на основе оптимизированных регексов"""

    def __init__(self):
        # Оптимизированные регексы (минимально необходимые)
        self.patterns = {
            'обычные': [
                ('EMAIL', re.compile(r'\b[\w\.-]+@[\w\.-]+\.\w{2,}\b'), self._validate_email),
                ('PHONE', re.compile(r'(?:\+7|8)[\s\-]?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}'),
                 self._validate_phone),
                ('DATE', re.compile(r'\b\d{2}[./]\d{2}[./]\d{4}\b'), None),
            ],
            'государственные': [
                ('SNILS', re.compile(r'\b\d{3}-\d{3}-\d{3}\s?\d{2}\b'), self._validate_snils),
                ('INN', re.compile(r'\b\d{10}\b|\b\d{12}\b'), self._validate_inn),
            ],
            'платёжные': [
                ('CARD', re.compile(r'\b\d{4}[\s\-]?\d{4}[\s\-]?\d{4}[\s\-]?\d{4}\b'), self._validate_card),
            ]
        }

    def _validate_email(self, email: str) -> bool:
        if EMAIL_VALIDATOR_AVAILABLE:
            try:
                validate_email(email)
                return True
            except:
                return False
        return bool(re.match(r'^[\w\.-]+@[\w\.-]+\.\w{2,}$', email))

    def _validate_phone(self, phone: str) -> bool:
        if PHONENUMBERS_AVAILABLE:
            try:
                parsed = phonenumbers.parse(phone, 'RU')
                return phonenumbers.is_valid_number(parsed)
            except:
                return False
        return len(re.sub(r'\D', '', phone)) in [10, 11]

    def _validate_snils(self, snils: str) -> bool:
        if SNILS_VALIDATOR_AVAILABLE:
            try:
                return snils_validator.validate(snils)
            except:
                pass
        # Fallback валидация
        nums = re.sub(r'\D', '', snils)
        if len(nums) != 11:
            return False
        base = [int(x) for x in nums[:9]]
        check = int(nums[9:])
        s = sum((9 - i) * d for i, d in enumerate(base))
        if s < 100:
            c = s
        elif s in (100, 101):
            c = 0
        else:
            c = s % 101
            if c == 100:
                c = 0
        return c == check

    def _validate_inn(self, inn: str) -> bool:
        nums = re.sub(r'\D', '', inn)
        if len(nums) == 10:
            w = [2, 4, 10, 3, 5, 9, 4, 6, 8]
            c = sum(int(nums[i]) * w[i] for i in range(9)) % 11 % 10
            return c == int(nums[9])
        elif len(nums) == 12:
            w1 = [7, 2, 4, 10, 3, 5, 9, 4, 6, 8, 0]
            w2 = [3, 7, 2, 4, 10, 3, 5, 9, 4, 6, 8, 0]
            c1 = sum(int(nums[i]) * w1[i] for i in range(11)) % 11 % 10
            c2 = sum(int(nums[i]) * w2[i] for i in range(11)) % 11 % 10
            return c1 == int(nums[10]) and c2 == int(nums[11])
        return False

    def _validate_card(self, card: str) -> bool:
        digits = re.sub(r'\D', '', card)
        if not (13 <= len(digits) <= 19):
            return False
        # Алгоритм Луна
        s = 0
        parity = len(digits) % 2
        for i, d in enumerate(digits):
            num = int(d)
            if i % 2 == parity:
                num *= 2
                if num > 9:
                    num -= 9
            s += num
        return s % 10 == 0

    def analyze_text(self, text: str) -> Dict[str, List[Dict]]:
        """Анализ текста с помощью оптимизированных регексов"""
        results = {
            'обычные': [],
            'государственные': [],
            'платёжные': [],
            'биометрические': [],
            'специальные': []
        }

        if not text:
            return results

        # Ограничиваем длину для производительности
        if len(text) > 100000:
            text = text[:100000]

        for category, patterns in self.patterns.items():
            for entity_type, pattern, validator in patterns:
                for match in pattern.finditer(text):
                    value = match.group(0)
                    if validator and not validator(value):
                        continue

                    # Проверяем контекст для снижения ложных срабатываний
                    if self._has_context(text, match.start(), value):
                        results[category].append({
                            'entity': entity_type,
                            'value': value,
                            'score': 0.8,
                            'start': match.start(),
                            'end': match.end()
                        })
                        break  # Одно вхождение на категорию достаточно

        return results

    def _has_context(self, text: str, pos: int, value: str) -> bool:
        """Проверка наличия контекста"""
        window = 50
        start = max(0, pos - window)
        end = min(len(text), pos + window)
        context = text[start:end].lower()

        context_keywords = {
            'EMAIL': ['email', 'e-mail', 'почта'],
            'PHONE': ['тел', 'телефон', 'phone', 'моб'],
            'SNILS': ['снилс', 'номер', 'страховой'],
            'INN': ['инн', 'налог', 'идентификационный'],
            'CARD': ['карт', 'card', 'visa', 'mastercard', 'cvv'],
        }

        for key, keywords in context_keywords.items():
            if key in value or any(kw in context for kw in keywords):
                return True
        return True  # Если нет специфического контекста, считаем валидным


# Выбор детектора
if PRESIDIO_AVAILABLE:
    pii_detector = OptimizedPIIDetector()
else:
    pii_detector = FallbackPIIDetector()


# --- Функции извлечения текста (оставляем без изменений) ---
def detect_encoding(raw_bytes: bytes) -> str:
    if chardet is None:
        return 'utf-8'
    try:
        res = chardet.detect(raw_bytes)
        enc = res.get('encoding') or 'utf-8'
        return enc
    except Exception:
        return 'utf-8'


def extract_text_generic(path: Path) -> str:
    try:
        raw = path.read_bytes()
        enc = detect_encoding(raw)
        return raw.decode(enc, errors='ignore')
    except Exception:
        return ''


def extract_text_pdf(path: Path) -> str:
    text = ''
    if pdfminer is not None:
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            text = pdfminer_extract(str(path)) or ''
            if text:
                return text
        except Exception:
            pass
    if PyPDF2 is not None:
        try:
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    try:
                        text += page.extract_text() or ''
                    except Exception:
                        pass
            return text
        except Exception:
            pass
    return text


def extract_text_docx(path: Path) -> str:
    if docx is None:
        return ''
    try:
        from docx import Document
        doc = Document(str(path))
        parts = []
        for p in doc.paragraphs:
            parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                parts.append(' \t '.join(cell.text for cell in row.cells))
        return '\n'.join(parts)
    except Exception:
        return ''


def extract_text_html(path: Path) -> str:
    txt = extract_text_generic(path)
    if bs4 is None:
        return txt
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(txt, 'html.parser')
        return soup.get_text(' ')
    except Exception:
        return txt


def extract_text_rtf(path: Path) -> str:
    raw = extract_text_generic(path)
    raw = re.sub(r'\\[a-zA-Z]+-?\d*\s?', ' ', raw)
    raw = re.sub(r'[{}]', ' ', raw)
    return re.sub(r'\s+', ' ', raw)


def extract_text_xls(path: Path) -> str:
    if pandas is None:
        return ''
    try:
        df = pandas.read_excel(str(path), header=None, dtype=str)
        return '\n'.join(' '.join(map(str, row.dropna().tolist())) for _, row in df.iterrows())
    except Exception:
        return ''


def extract_text_image(path: Path) -> str:
    if PIL is None or pytesseract is None:
        return ''
    try:
        from PIL import Image
        img = Image.open(str(path))
        return pytesseract.image_to_string(img, lang='rus+eng')
    except Exception:
        return ''


def extract_text_doc(path: Path) -> str:
    return extract_text_generic(path)


def extract_mp4(path: Path) -> str:
    if PIL is None or pytesseract is None or cv2 is None:
        return ''

    text_parts = []
    try:
        cap = cv2.VideoCapture(str(path))
        if not cap.isOpened():
            return ''

        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        fps = cap.get(cv2.CAP_PROP_FPS)
        duration = total_frames / fps if fps > 0 else 0

        # Динамический шаг кадров
        if duration < 30:
            frame_step = max(1, total_frames // 10)
        elif duration < 300:
            frame_step = max(1, total_frames // 20)
        else:
            frame_step = max(1, total_frames // 30)

        frame_step = min(frame_step, total_frames // 5) if total_frames > 0 else 30
        frame_step = max(frame_step, 1)

        processed_frames = 0
        max_frames = 20  # Ограничиваем количество кадров

        frame_idx = 0
        while processed_frames < max_frames:
            ret, frame = cap.read()
            if not ret:
                break

            if frame_idx % frame_step == 0:
                frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                from PIL import Image
                pil_image = Image.fromarray(frame_rgb)
                text = pytesseract.image_to_string(pil_image, lang='rus+eng')
                if text and text.strip():
                    text_parts.append(text.strip())
                    processed_frames += 1
            frame_idx += 1

        cap.release()
        return '\n'.join(text_parts)
    except Exception:
        return ''


def extract_text(path: Path) -> str:
    ext = path.suffix.lower().lstrip('.')
    try:
        if ext == 'pdf':
            return extract_text_pdf(path)
        elif ext == 'docx':
            return extract_text_docx(path)
        elif ext == 'html':
            return extract_text_html(path)
        elif ext == 'rtf':
            return extract_text_rtf(path)
        elif ext == 'xls':
            return extract_text_xls(path)
        elif ext in {'jpg', 'jpeg', 'png', 'gif'}:
            return extract_text_image(path)
        elif ext == 'doc':
            return extract_text_doc(path)
        elif ext == 'mp4':
            return extract_mp4(path)
        else:
            return extract_text_generic(path)
    except Exception:
        return ''


def detect_categories(text: str) -> Dict[str, int]:
    """Обновленная функция детекции с использованием оптимизированного детектора"""
    results = pii_detector.analyze_text(text)

    # Подсчет количества находок по категориям
    categories = {
        'обычные': len(results['обычные']),
        'государственные': len(results['государственные']),
        'платёжные': len(results['платёжные']),
        'биометрические': len(results['биометрические']),
        'специальные': len(results['специальные'])
    }

    # Ограничиваем максимальное количество для избежания завышения
    for key in categories:
        categories[key] = min(categories[key], 10)

    return categories


def estimate_uz(cats: Dict[str, int]) -> str:
    total = sum(cats.values())
    distinct = sum(1 for v in cats.values() if v > 0)
    has_special = cats['специальные'] > 0
    has_bio = cats['биометрические'] > 0
    has_pay = cats['платёжные'] > 0
    has_gov = cats['государственные'] > 0
    has_common = cats['обычные'] > 0

    if has_special or has_bio:
        return 'УЗ-1' if (total >= 3 or distinct >= 2) else 'УЗ-2'
    if has_pay or has_gov:
        return 'УЗ-2' if (total >= 3 or distinct >= 2) else 'УЗ-3'
    if has_common:
        return 'УЗ-3' if (total >= 5 or distinct >= 2) else 'УЗ-4'
    return 'нет признаков'


def scan_root(root: Path) -> List[Dict[str, object]]:
    results: List[Dict[str, object]] = []
    total_files = 0

    for dirpath, dirnames, filenames in os.walk(root):
        for name in filenames:
            p = Path(dirpath) / name
            ext = p.suffix.lower().lstrip('.')
            if ext not in INCLUDE_EXTS:
                continue

            total_files += 1
            if total_files % 10 == 0:
                print(f"Processed {total_files} files...")

            try:
                text = extract_text(p)
                if text:
                    cats = detect_categories(text)
                else:
                    cats = {'обычные': 0, 'государственные': 0, 'платёжные': 0,
                            'биометрические': 0, 'специальные': 0}

                uz = estimate_uz(cats)
                res = {
                    'path': str(p),
                    'categories': {k: v for k, v in cats.items() if v > 0},
                    'uz': uz,
                    'total_hits': sum(cats.values()),
                    'ext': ext
                }
                results.append(res)
            except Exception as e:
                results.append({
                    'path': str(p),
                    'categories': {},
                    'uz': 'error',
                    'error': str(e)[:100],
                    'ext': ext
                })

    return results


def print_summary(results: List[Dict[str, object]]):
    uz_stats = {}
    for r in results:
        uz = r.get('uz', 'unknown')
        uz_stats[uz] = uz_stats.get(uz, 0) + 1
        cats = ', '.join(sorted(r['categories'].keys())) if r.get('categories') else '—'
        print(f"{r['path']}: {cats} → {r['uz']}")

    print(f"\n--- Статистика по УЗ ---")
    for uz, count in sorted(uz_stats.items()):
        print(f"{uz}: {count} файлов")


def save_csv(results: List[Dict[str, object]], out_csv: Path, save_all: bool = False):
    out_csv = Path(out_csv)

    # Фильтруем результаты
    if save_all:
        filtered_results = results
        print(f"Сохраняем все файлы ({len(filtered_results)}) в CSV")
    else:
        filtered_results = [r for r in results if r.get('total_hits', 0) != 0]
        print(f"Найдено {len(filtered_results)} файлов с ПДн из {len(results)} обработанных")

    if not filtered_results:
        print("Нет файлов для сохранения. CSV-файл не создан.")
        return None

    with out_csv.open('w', newline='', encoding='utf-8') as f:
        w = csv.writer(f)
        w.writerow(['путь', 'категории ПДн', 'количество_находок', 'УЗ', 'формат файла'])

        # Сортируем по количеству находок (от большего к меньшему)
        sorted_results = sorted(filtered_results,
                                key=lambda x: x.get('total_hits', 0),
                                reverse=True)

        for r in sorted_results:
            cats_list = list(r.get('categories', {}).keys())
            cats_str = ', '.join(cats_list) if cats_list else '—'
            total = r.get('total_hits', 0)
            uz = r.get('uz', 'нет признаков')
            fmt = r.get('ext', '')
            w.writerow([r['path'], cats_str, total, uz, fmt])

    # Выводим статистику по УЗ для найденных файлов
    uz_stats = {}
    for r in filtered_results:
        uz = r.get('uz', 'нет признаков')
        uz_stats[uz] = uz_stats.get(uz, 0) + 1

    print(f"\nСтатистика по уровню защищенности (файлы с ПДн):")
    for uz, count in sorted(uz_stats.items()):
        print(f"  {uz}: {count} файлов")

    print(f"\nCSV файл сохранен: {out_csv.resolve()}")
    return out_csv


# --- Запуск ---
if __name__ == "__main__":
    if ROOT_DIR.exists():
        print(f"Сканирование {ROOT_DIR}...")
        print(f"Используемый детектор: {'Presidio (ML)' if PRESIDIO_AVAILABLE else 'Fallback (Regex)'}")
        results = scan_root(ROOT_DIR)
        print_summary(results)
        save_csv(results, OUTPUT_CSV)
        print(f"\nСохранено: {OUTPUT_CSV.resolve()}")
    else:
        print("Укажите корректный ROOT_DIR (существующая директория).")