from __future__ import annotations
import os, re, io, json, math, csv, itertools, mimetypes
from pathlib import Path
from typing import Dict, List, Tuple, Optional

ROOT_DIR = Path('D:/dataset/share')  # ← укажите корень сканирования
OUTPUT_CSV = Path('pii_scan_results.csv')
INCLUDE_EXTS = {'doc','docx','gif','html','ipynb','jpeg','jpg','pdf','php','png','rtf','xls'}

def safe_import(name):
    try:
        return __import__(name)
    except Exception:
        return None

PyPDF2 = safe_import('PyPDF2')
pdfminer = safe_import('pdfminer')
docx = safe_import('docx')
bs4 = safe_import('bs4')
pandas = safe_import('pandas')
PIL = safe_import('PIL')
pytesseract = safe_import('pytesseract')
chardet = safe_import('chardet')

def detect_encoding(raw_bytes: bytes) -> str:
    if chardet is None:
        return 'utf-8'
    try:
        res = chardet.detect(raw_bytes)
        enc = res.get('encoding') or 'utf-8'
        return enc
    except Exception:
        return 'utf-8'
# --- Вспомогательные валидаторы и функции ---
def luhn_check(number: str) -> bool:
    digits = [int(d) for d in re.sub(r'\D', '', number)]
    if not (13 <= len(digits) <= 19):
        return False
    s = 0
    parity = len(digits) % 2
    for i, d in enumerate(digits):
        if i % 2 == parity:
            d *= 2
            if d > 9:
                d -= 9
        s += d
    return s % 10 == 0

def snils_valid(snils: str) -> bool:
    nums = re.sub(r'\D', '', snils)
    if len(nums) != 11:
        return False
    base = [int(x) for x in nums[:9]]
    check = int(nums[9:])
    s = sum((9 - i) * d for i, d in enumerate(base))
    if s < 100:
        c = s
    elif s in (100, 101):
        c = 0
    else:
        c = s % 101
        if c == 100:
            c = 0
    return c == check

def inn_valid(inn: str) -> bool:
    nums = re.sub(r'\D', '', inn)
    if len(nums) == 10:
        w = [2,4,10,3,5,9,4,6,8]
        c = sum(int(nums[i]) * w[i] for i in range(9)) % 11 % 10
        return c == int(nums[9])
    elif len(nums) == 12:
        w1 = [7,2,4,10,3,5,9,4,6,8,0]
        w2 = [3,7,2,4,10,3,5,9,4,6,8,0]
        c1 = sum(int(nums[i]) * w1[i] for i in range(11)) % 11 % 10
        c2 = sum(int(nums[i]) * w2[i] for i in range(11)) % 11 % 10
        return c1 == int(nums[10]) and c2 == int(nums[11])
    return False

def has_context(text: str, idx: int, window: int, *keywords: str) -> bool:
    start = max(0, idx - window)
    end = min(len(text), idx + window)
    chunk = text[start:end]
    return any(k in chunk for k in keywords)
# --- Извлечение текста по типам ---
def extract_text_generic(path: Path) -> str:
    # попытка прочитать как текст с автоопределением кодировки
    try:
        raw = path.read_bytes()
        enc = detect_encoding(raw)
        return raw.decode(enc, errors='ignore')
    except Exception:
        return ''

def extract_text_pdf(path: Path) -> str:
    text = ''
    if pdfminer is not None:
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            text = pdfminer_extract(str(path)) or ''
            if text:
                return text
        except Exception:
            pass
    if PyPDF2 is not None:
        try:
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    try:
                        text += page.extract_text() or ''
                    except Exception:
                        pass
            return text
        except Exception:
            pass
    return text

def extract_text_docx(path: Path) -> str:
    if docx is None:
        return ''
    try:
        from docx import Document
        doc = Document(str(path))
        parts = []
        for p in doc.paragraphs:
            parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                parts.append(' \t '.join(cell.text for cell in row.cells))
        return '\n'.join(parts)
    except Exception:
        return ''

def extract_text_html(path: Path) -> str:
    txt = extract_text_generic(path)
    if bs4 is None:
        return txt
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(txt, 'lxml') if 'lxml' in str(bs4.builder.__dict__) else BeautifulSoup(txt, 'html.parser')
        return soup.get_text(' ')
    except Exception:
        return txt

def extract_text_rtf(path: Path) -> str:
    raw = extract_text_generic(path)
    # грубое снятие управляющих последовательностей RTF
    raw = re.sub(r'\\[a-zA-Z]+-?\d*\s?', ' ', raw)
    raw = re.sub(r'[{}]', ' ', raw)
    return re.sub(r'\s+', ' ', raw)

def extract_text_ipynb(path: Path) -> str:
    try:
        data = json.loads(path.read_text(encoding='utf-8', errors='ignore'))
        parts = []
        for cell in data.get('cells', []):
            src = cell.get('source', [])
            if isinstance(src, list):
                parts.append(''.join(src))
            elif isinstance(src, str):
                parts.append(src)
        return '\n'.join(parts)
    except Exception:
        return ''

def extract_text_xls(path: Path) -> str:
    if pandas is None:
        return ''
    try:
        # для .xls нужен xlrd; для .xlsx (если попадётся) — openpyxl
        df = pandas.read_excel(str(path), header=None, dtype=str)
        return '\n'.join(' '.join(map(str, row.dropna().tolist())) for _, row in df.iterrows())
    except Exception:
        return ''

def extract_text_image(path: Path) -> str:
    if PIL is None or pytesseract is None:
        return ''
    try:
        from PIL import Image
        img = Image.open(str(path))
        return pytesseract.image_to_string(img, lang='rus+eng')
    except Exception:
        return ''

def extract_text_doc(path: Path) -> str:
    # .doc (старый формат) без внешних утилит разобрать сложно; оставим как бинарный текстовый скимминг
    raw = extract_text_generic(path)
    return raw

def extract_text(path: Path) -> str:
    ext = path.suffix.lower().lstrip('.')
    try:
        if ext == 'pdf':
            return extract_text_pdf(path)
        elif ext == 'docx':
            return extract_text_docx(path)
        elif ext in {'html','php'}:
            return extract_text_html(path)
        elif ext == 'rtf':
            return extract_text_rtf(path)
        elif ext == 'ipynb':
            return extract_text_ipynb(path)
        elif ext == 'xls':
            return extract_text_xls(path)
        elif ext in {'jpg','jpeg','png','gif'}:
            return extract_text_image(path)
        elif ext == 'doc':
            return extract_text_doc(path)
        else:
            return extract_text_generic(path)
    except Exception:
        return ''
# --- Детекторы ПДн (регексы и ключевые слова) ---
EMAIL_RE = re.compile(r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[A-Za-z]{2,}\b")
PHONE_RE = re.compile(r"(?:(?:\+7|8)\s*\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{2}[\s-]?\d{2})")
FIO_RE = re.compile(r"\b[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+)?\b")
DOB_RE = re.compile(r"\b(\d{2}[./]\d{2}[./]\d{4})\b")
INDEX_RE = re.compile(r"\b\d{6}\b")

# Гос. идентификаторы
SNILS_RE = re.compile(r"\b\d{3}-\d{3}-\d{3}\s?\d{2}\b")
INN10_RE = re.compile(r"(?<!\d)\d{10}(?!\d)")
INN12_RE = re.compile(r"(?<!\d)\d{12}(?!\d)")
PASSPORT_RE = re.compile(r"(?:(?<!\d)\d{2}\s?\d{2}\s?\d{6}(?!\d))")
MRZ_RE = re.compile(r"[P|V|C]<[A-Z<]{2}")
DL_RE = re.compile(r"(?<!\d)\d{10,12}(?!\d)")  # водительское удостоверение (грубая форма)

# Платёжные
CARD_RE = re.compile(r"(?:(?:\d[ -]*?){13,19})")
CVV_RE = re.compile(r"\b(CVV|CVC|CVV2)\b", re.IGNORECASE)
RS_RE = re.compile(r"(?i)(?:р/с|расч[её]тн(?:ый)?\s+сч[её]т)[^\d]*(\d{20})")
BIK_RE = re.compile(r"(?i)бик[^\d]*(\d{9})")

# Биометрия/специальные — по ключевым словам
BIOMETRIC_KEYS = [
    'биометр', 'отпечат', 'радуж', 'ирис', 'лицев', 'селфи', 'faceid', 'fingerprint', 'iris', 'voiceprint', 'голосов', 'геометрия лица'
]
SPECIAL_KEYS = [
    'диагноз', 'анамнез', 'инвалид', 'здоровь', 'медицин', 'психиатр', 'вич', 'религ', 'вероисповед', 'политическ', 'партия', 'интим', 'сексуаль'
]

def count_occurrences(pattern: re.Pattern, text: str) -> int:
    return len(list(pattern.finditer(text)))

def detect_categories(text: str) -> Dict[str, int]:
    t = text if isinstance(text, str) else ''
    low = t.lower()
    cats = { 'обычные': 0, 'государственные': 0, 'платёжные': 0, 'биометрические': 0, 'специальные': 0 }

    # Обычные
    cats['обычные'] += count_occurrences(EMAIL_RE, t)
    cats['обычные'] += count_occurrences(PHONE_RE, t)
    # ФИО — не слишком агрессивно: ограничим, чтобы не раздувать
    cats['обычные'] += min(5, count_occurrences(FIO_RE, t))
    # ДР чаще засчитываем только при наличии контекста
    for m in DOB_RE.finditer(t):
        if has_context(low, m.start(), 40, 'дата рождения', 'родил'):
            cats['обычные'] += 1
    # Адрес — индекс + ключи
    for m in INDEX_RE.finditer(t):
        if has_context(low, m.start(), 40, 'ул', 'улица', 'просп', 'пер', 'дом', 'квартира', 'город', 'г.'):
            cats['обычные'] += 1

    # Государственные идентификаторы
    for m in SNILS_RE.finditer(t):
        if snils_valid(m.group(0)):
            cats['государственные'] += 1
    for m in INN10_RE.finditer(t):
        s = m.group(0)
        if inn_valid(s):
            cats['государственные'] += 1
    for m in INN12_RE.finditer(t):
        s = m.group(0)
        if inn_valid(s):
            cats['государственные'] += 1
    for m in PASSPORT_RE.finditer(t):
        if has_context(low, m.start(), 50, 'паспорт', 'серия', 'номер', 'код подразделения'):
            cats['государственные'] += 1
    for m in DL_RE.finditer(t):
        if has_context(low, m.start(), 30, 'водител', 'удостовер'):
            cats['государственные'] += 1
    if MRZ_RE.search(t):
        cats['государственные'] += 1

    # Платёжные
    for m in CARD_RE.finditer(t):
        raw = m.group(0)
        digits = re.sub(r'\D', '', raw)
        if 13 <= len(digits) <= 19 and luhn_check(raw):
            # требуем контекст, чтобы снизить FP
            if has_context(low, m.start(), 40, 'visa', 'mastercard', 'карта', 'cvv', 'cvc', 'номер карты'):
                cats['платёжные'] += 1
    for m in RS_RE.finditer(t):
        cats['платёжные'] += 1
    for m in BIK_RE.finditer(t):
        cats['платёжные'] += 1
    if CVV_RE.search(t):
        cats['платёжные'] += 1

    # Биометрические
    if any(k in low for k in BIOMETRIC_KEYS):
        cats['биометрические'] += 1

    # Специальные (ключевые слова)
    if any(k in low for k in SPECIAL_KEYS):
        cats['специальные'] += 1

    return cats

def estimate_uz(cats: Dict[str, int]) -> str:
    total = sum(cats.values())
    distinct = sum(1 for v in cats.values() if v > 0)
    has_special = cats['специальные'] > 0
    has_bio = cats['биометрические'] > 0
    has_pay = cats['платёжные'] > 0
    has_gov = cats['государственные'] > 0
    has_common = cats['обычные'] > 0

    if has_special or has_bio:
        return 'УЗ-1' if (total >= 5 or distinct >= 2) else 'УЗ-2'
    if has_pay or has_gov:
        return 'УЗ-2' if (total >= 5 or distinct >= 2) else 'УЗ-3'
    if has_common:
        return 'УЗ-3' if (total >= 5 or distinct >= 2) else 'УЗ-4'
    return 'нет признаков'
# --- Обход папки, детекция, агрегированный вывод ---
from datetime import datetime

def scan_root(root: Path) -> List[Dict[str, object]]:
    results: List[Dict[str, object]] = []
    for dirpath, dirnames, filenames in os.walk(root):
        for name in filenames:
            p = Path(dirpath) / name
            ext = p.suffix.lower().lstrip('.')
            if ext not in INCLUDE_EXTS:
                continue
            try:
                text = extract_text(p)
                cats = detect_categories(text)
                uz = estimate_uz(cats)
                res = {
                    'path': str(p),
                    'categories': {k:v for k,v in cats.items() if v>0},
                    'uz': uz,
                    'total_hits': sum(cats.values()),
                    'ext': ext
                }
                results.append(res)
            except Exception as e:
                results.append({'path': str(p), 'categories': {}, 'uz': 'error', 'error': str(e), 'ext': ext})
    return results

def print_summary(results: List[Dict[str, object]]):
    for r in results:
        cats = ', '.join(sorted(r['categories'].keys())) if r.get('categories') else '—'
        print(f"{r['path']}: {cats} → {r['uz']}")

def save_csv(results: List[Dict[str, object]], out_csv: Path):
    out_csv = Path(out_csv)
    with out_csv.open('w', newline='', encoding='utf-8') as f:
        w = csv.writer(f)
        w.writerow(['path','categories','uz','total_hits','ext'])
        for r in results:
            w.writerow([r['path'], json.dumps(r['categories'], ensure_ascii=False), r['uz'], r.get('total_hits',0), r.get('ext','')])
    return out_csv

# --- Запуск ---
if ROOT_DIR.exists():
    results = scan_root(ROOT_DIR)
    print_summary(results)
    save_csv(results, OUTPUT_CSV)
    print(f"\nСохранено: {OUTPUT_CSV.resolve()}")
else:
    print("Укажите корректный ROOT_DIR (существующая директория).")
